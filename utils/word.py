from docx import Document
from io import BytesIO

def create_passport_translation_doc(fields: dict) -> bytes:
    doc = Document()
    def up(s):
        return s.upper() if isinstance(s, str) else s
    def region_ru(s):
        return s.replace('REGION', 'ОБЛАСТЬ').replace('region', 'ОБЛАСТЬ') if isinstance(s, str) else s

    doc.add_paragraph("Перевод выполнен с узбекского и английского языков на русский язык/")
    doc.add_paragraph("")
    doc.add_paragraph("РЕСПУБЛИКА УЗБЕКИСТАН\nРЕСПУБЛИКА УЗБЕКИСТАН\nУЗБЕКИСТАН\nПАСПОРТ\nПАСПОРТ\n{}".format(up(fields.get('passport_number', ''))))
    doc.add_paragraph("")
    doc.add_paragraph("РЕСПУБЛИКА УЗБЕКИСТАН / РЕСПУБЛИКА УЗБЕКИСТАН")
    # Таблица как в паспорте: ПАСПОРТ | ТИП | КОД СТРАНЫ | НОМЕР ПАСПОРТА
    table = doc.add_table(rows=2, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = up("ПАСПОРТ")
    hdr[1].text = up("ТИП")
    hdr[2].text = up("КОД СТРАНЫ")
    hdr[3].text = up("НОМЕР ПАСПОРТА")
    row = table.rows[1].cells
    row[0].text = up("ПАСПОРТ")
    row[1].text = up(fields.get('passport_type', 'P'))
    row[2].text = up(fields.get('country_code', 'UZB'))
    row[3].text = up(fields.get('passport_number', ''))

    # Блок ФИО: Фамилия | Имя | Отчество
    fio_table = doc.add_table(rows=1, cols=3)
    fio_hdr = fio_table.rows[0].cells
    fio_hdr[0].text = up("ФАМИЛИЯ")
    fio_hdr[1].text = up("ИМЯ")
    fio_hdr[2].text = up("ОТЧЕСТВО")
    fio_row = fio_table.add_row().cells
    # Разделяем ФИО по пробелам (ожидается: Фамилия Имя Отчество)
    fio_parts = fields.get('fio', '').split()
    fio_row[0].text = up(fio_parts[0]) if len(fio_parts) > 0 else ""
    fio_row[1].text = up(fio_parts[1]) if len(fio_parts) > 1 else ""
    fio_row[2].text = up(' '.join(fio_parts[2:])) if len(fio_parts) > 2 else ""


    # Далее остальные данные
    def add_line(label, value):
        doc.add_paragraph(f"{up(label)}    {up(region_ru(value))}")

    add_line("ГРАЖДАНСТВО", fields.get('nationality', ''))
    add_line("ДАТА РОЖДЕНИЯ", fields.get('birthdate', ''))
    add_line("МЕСТО РОЖДЕНИЯ", fields.get('birth_place', ''))
    add_line("ПОЛ", fields.get('sex', ''))
    add_line("ДАТА ВЫДАЧИ", fields.get('issue_date', ''))
    add_line("ДЕЙСТВИТЕЛЕН ДО", fields.get('expiry_date', ''))
    add_line("ОРГАН ВЫДАЧИ", fields.get('authority', ''))

    doc.add_paragraph("")
    doc.add_paragraph("Стр.3\nРЕСПУБЛИКА УЗБЕКИСТАН\nРЕСПУБЛИКА УЗБЕКИСТАН\nUZB\n/подписано/\nподпись владельца\n{}".format(up(fields.get('passport_number', ''))))
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
