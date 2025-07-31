import logging
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def fill_table_row_char_by_char(table, text):
    """Заполняет ячейки первой строки таблицы по одному символу."""
    if not table.rows:
        logger.warning("Попытка заполнить таблицу без строк.")
        return
    
    row = table.rows[0]
    # Очищаем все ячейки перед заполнением, чтобы избежать дублирования
    for cell in row.cells:
        if cell.paragraphs:
            p = cell.paragraphs[0]
            p.clear()

    # Заполняем ячейки по одному символу
    for i, char in enumerate(text):
        if i < len(row.cells):
            cell = row.cells[i]
            # Убедимся, что параграф существует
            if not cell.paragraphs:
                p = cell.add_paragraph()
            else:
                p = cell.paragraphs[0]
            
            # Добавляем новый символ
            run = p.add_run(char)
            
            # Устанавливаем форматирование
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Если символов больше, чем ячеек, прекращаем заполнение
            logger.warning(f"Текст '{text}' слишком длинный для таблицы. Заполнено {i} символов.")
            break

def generate_notification_word(data, template_path, output_path):
    """
    Генерирует уведомление в формате Word, заполняя шаблон данными.
    Каждый символ данных помещается в отдельную ячейку таблицы.
    """
    try:
        document = Document(template_path)
        tables = document.tables
        num_tables = len(tables)
        logger.info(f"Шаблон '{template_path}' успешно открыт. Количество таблиц: {num_tables}")

        # Карта соответствия данных и ИНДЕКСОВ таблиц (от 0 до 19)
        fill_map = {
            0: data.get('recipient_department', ''),
            1: data.get('last_name', ''),
            2: data.get('first_name', ''),
            3: data.get('middle_name', ''),
            4: data.get('citizenship', ''),
            5: data.get('birth_date', ''),
            6: data.get('passport_series', ''),
            7: data.get('passport_number', ''),
            8: data.get('passport_issuer', ''),
            9: data.get('passport_issue_date', ''),
            10: data.get('patent_series', ''),
            11: data.get('patent_number', ''),
            12: data.get('patent_issue_date', ''),
            13: data.get('profession', ''),
            14: data.get('work_address', ''),
            15: data.get('contract_date', ''),
            16: data.get('inn', ''),
            17: data.get('insurance_policy_details', ''),
            18: data.get('insurance_policy_series', ''),
            19: data.get('insurance_policy_issue_date', ''),
        }

        # Заполняем таблицы, строго проверяя границы
        for table_index, text_to_fill in fill_map.items():
            # Проверяем, есть ли данные для заполнения
            if text_to_fill:
                # Главная проверка: убедимся, что индекс не выходит за пределы
                if table_index < num_tables:
                    logger.info(f"Заполнение таблицы {table_index} данными: '{text_to_fill}'")
                    fill_table_row_char_by_char(tables[table_index], str(text_to_fill))
                else:
                    logger.error(f"Критическая ошибка: Попытка доступа к таблице {table_index}, но в документе всего {num_tables} таблиц.")
                    # Можно прервать выполнение или просто пропустить эту таблицу
                    continue
        
        # Дополнительно можно обработать поля, которые не в таблицах, если они есть
        # Например, дата в самом низу документа
        # paragraph = document.paragraphs[-1] # Пример: последний параграф
        # paragraph.text = paragraph.text.replace("20__ г.", f"20{data.get('year', '__')} г.")


        document.save(output_path)
        logger.info(f"Уведомление успешно создано и сохранено в {output_path}")
        return output_path

    except FileNotFoundError:
        logger.error(f"Ошибка: Файл шаблона не найден по пути: {template_path}")
        return None
    except Exception as e:
        logger.error(f"Произошла непредвиденная ошибка при создании документа: {e}", exc_info=True)
        return None
